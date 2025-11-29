import pdfplumber
from docx import Document

def extract_text_from_pdf(path: str):
    """Extract text from PDF using pdfplumber."""
    text = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                # improve extraction accuracy
                t = page.extract_text(x_tolerance=1)
                if t:
                    text.append(t)
    except Exception as e:
        return f"ERROR reading PDF: {e}"

    return "\n".join(text) if text else " "

def extract_text_from_docx(path: str):
    """Extract text from DOCX including paragraphs & tables."""
    try:
        doc = Document(path)
        parts = []

        # paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        # tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        return "\n".join(parts) if parts else " "

    except Exception as e:
        return f"ERROR reading DOCX: {e}"

def extract_text(path: str):
    """
    Detect file type WITHOUT altering path string.
    """
    lower = path.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(path)
    elif lower.endswith(".docx"):
        return extract_text_from_docx(path)
    else:
        return " "
